"""
Extração de texto de arquivos PDF e DOCX.
Retorna o texto limpo como string.

PDFs com tabelas
----------------
PDFs frequentemente carregam dados em grades (grade curricular, plano
de ensino, ementas). Extração ingênua via ``get_text`` cospe linhas
tipo ``"ADS1 Algoritmos 120 - 1"`` — o cabeçalho da coluna some, o
chunker corta no meio e o retriever nunca casa "qual a carga horária
de AEDs1?" com "120h".

Estratégia: ``page.find_tables()`` detecta grades; cada linha vira uma
sentença ``"Disciplina: AEDs1; CH: 120; Período: 1"`` que sobrevive
ao chunking porque o cabeçalho é repetido em CADA linha. Os blocos de
texto que caem dentro do bbox da tabela são pulados na extração normal
pra não duplicar (e poluir) o índice com a versão fragmentada.
"""

import io
import re
import unicodedata

import fitz          # PyMuPDF
import docx          # python-docx


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Detecta o tipo pelo nome do arquivo e extrai o texto.
    Levanta ValueError para tipos não suportados.
    """
    name_lower = filename.lower()

    if name_lower.endswith(".pdf"):
        return _extract_pdf(file_bytes)
    elif name_lower.endswith(".docx"):
        return _extract_docx(file_bytes)
    elif name_lower.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"Tipo de arquivo não suportado: {filename}")


def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extrai texto de PDF usando PyMuPDF (fitz).
    Preserva a estrutura de parágrafos com quebras de linha.
    Tabelas detectadas são renderizadas como sentenças linha-a-linha
    com cabeçalho repetido (ver docstring do módulo).
    """
    text_parts = []

    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page_num, page in enumerate(doc, start=1):
            page_block = _extract_page(page, page_num)
            if page_block:
                text_parts.append(page_block)

    full_text = "\n\n".join(text_parts)
    return _clean_text(full_text)


def _extract_page(page, page_num: int) -> str:
    """Extrai uma página combinando texto não-tabular + tabelas renderizadas."""
    # 1. Detecta tabelas (best-effort — pode falhar em PDFs malformados)
    try:
        finder = page.find_tables()
        tables = list(finder) if finder else []
    except Exception:
        tables = []

    table_bboxes = [fitz.Rect(t.bbox) for t in tables]

    # 2. Texto não-tabular: pega blocos e descarta os que caem dentro de tabela
    non_table_text = _extract_non_table_blocks(page, table_bboxes)

    # 3. Renderiza cada tabela como sentenças com cabeçalho repetido
    table_sections = []
    for table in tables:
        rendered = _render_table_as_sentences(table)
        if rendered:
            table_sections.append(rendered)

    sections = []
    if non_table_text.strip():
        sections.append(non_table_text)
    sections.extend(table_sections)

    if not sections:
        return ""

    return f"[Página {page_num}]\n" + "\n\n".join(sections)


def _extract_non_table_blocks(page, table_bboxes: list) -> str:
    """Junta blocos de texto que NÃO estão dentro de uma tabela detectada.

    Sem isso, o conteúdo da tabela apareceria duas vezes: uma fragmentada
    (como blocos de texto soltos) e outra estruturada (como sentenças
    linha-a-linha). O par dilui o sinal no retrieval.
    """
    blocks = page.get_text("blocks")
    if not blocks:
        return ""

    parts = []
    for block in blocks:
        # block: (x0, y0, x1, y1, text, block_no, block_type)
        if len(block) < 5:
            continue
        block_rect = fitz.Rect(block[0], block[1], block[2], block[3])
        block_text = block[4]
        if not block_text or not block_text.strip():
            continue
        # Pula imagens (block_type == 1) e blocos que caem em tabela
        if len(block) >= 7 and block[6] != 0:
            continue
        if any(_rect_overlaps(block_rect, tb) for tb in table_bboxes):
            continue
        parts.append(block_text)

    return "\n".join(parts)


def _rect_overlaps(a: "fitz.Rect", b: "fitz.Rect", threshold: float = 0.5) -> bool:
    """True se mais de ``threshold`` da área de ``a`` está dentro de ``b``.

    Usamos a fração de ``a`` (não a interseção bruta) pra não descartar
    parágrafos grandes que só esbarram de raspão em uma tabela pequena.
    """
    inter = a & b
    if inter.is_empty:
        return False
    area_a = a.get_area()
    if area_a <= 0:
        return False
    return (inter.get_area() / area_a) >= threshold


def _render_table_as_sentences(table) -> str:
    """Converte uma tabela detectada em sentenças linha-a-linha.

    Cada linha de dados vira uma string ``"col_a: val_a; col_b: val_b."``,
    com o cabeçalho repetido em cada linha. Isso garante que mesmo se o
    chunker partir a tabela ao meio, todo chunk carrega o contexto das
    colunas — o que destrava queries do tipo "qual a CH de X?".
    """
    try:
        rows = table.extract()
    except Exception:
        return ""

    if not rows or len(rows) < 2:
        return ""

    headers, data_start = _extract_headers(rows)
    if not headers:
        return ""

    sentences = []
    for row in rows[data_start:]:
        cells = [_clean_cell(c) for c in row]
        if not any(cells):  # linha totalmente vazia
            continue
        pairs = []
        for h, v in zip(headers, cells):
            if v:
                pairs.append(f"{h}: {v}")
        if pairs:
            sentences.append("- " + "; ".join(pairs) + ".")

    # Une linhas com \n\n (e não \n) pra o chunker recursivo
    # PRIORIZAR quebrar entre linhas da tabela. Antes ele podia cortar
    # no meio de "Pré-Req: AEDs1; ID: AEDs2." perdendo a sigla AEDs2 —
    # vi isso acontecer com a Grade Computação. Com \n\n o splitter
    # mantém cada linha íntegra.
    return "\n\n".join(sentences)


def _extract_headers(rows: list) -> tuple[list[str], int]:
    """Identifica o cabeçalho da tabela e o índice da 1ª linha de dados.

    Retorna ``(headers, data_start_index)``.

    Caso comum (cabeçalho normal): ``row[0]`` tem pelo menos metade das
    colunas preenchidas — devolve esses valores como headers e
    ``data_start = 1``.

    Caso "banner row": ``row[0]`` tem só **1 célula preenchida** e as
    demais ``None``. É o padrão de PDFs que põem um título-faixa em cima
    da grade (ex.: "Grade Computação" com headers reais embutidos na
    mesma célula, separados por ``\\n``). Nesses casos:

    1. Tenta extrair os cabeçalhos reais da última linha do banner
       (heurística ``_parse_header_line`` — split por whitespace + merge
       em volta de ``/`` e ``-`` pra compor "Pré-Req / Co-req" como
       1 token).
    2. Se a heurística não conseguir bater o número de colunas, cai pro
       fallback genérico ``col_1..col_N`` (descartando o banner — o
       título do doc já vem do filename via metadata do Chroma).
    """
    if not rows:
        return [], 0

    raw_first = [_clean_cell_keep_newlines(c) for c in rows[0]]
    n_cols = len(raw_first)
    if n_cols == 0:
        return [], 1

    filled = sum(1 for c in raw_first if c)

    # Caso normal: ≥ metade das colunas preenchidas → row[0] é o header.
    if filled >= max(2, (n_cols + 1) // 2):
        headers = [
            _flat(h) if h else f"col_{i + 1}"
            for i, h in enumerate(raw_first)
        ]
        return headers, 1

    # Caso banner: 1 célula preenchida com o resto vazio. Tenta extrair
    # os headers reais da última linha do conteúdo do banner.
    if filled == 1 and raw_first[0]:
        banner_lines = [
            ln.strip() for ln in raw_first[0].split("\n") if ln.strip()
        ]
        if len(banner_lines) >= 2:
            parsed = _parse_header_line(banner_lines[-1], n_cols)
            if parsed:
                return parsed, 1

    # Fallback: nenhum header confiável — usa col_N e descarta row[0].
    return [f"col_{i + 1}" for i in range(n_cols)], 1


def _parse_header_line(line: str, n_cols: int) -> list[str] | None:
    """Tenta dividir uma linha de cabeçalho em ``n_cols`` tokens.

    Estratégia em camadas:

    1. Split por whitespace. Se já bater com ``n_cols``, devolve direto.
    2. Mescla tokens em volta de ``/`` (composições tipo
       ``"Pré-Requisito / Co-requisito"``).
    3. Mescla em volta de ``-`` isolado.
    4. Se nenhuma camada bater, devolve ``None`` — caller cai no
       fallback ``col_1..col_N``.
    """
    tokens = line.split()
    if not tokens:
        return None
    if len(tokens) == n_cols:
        return tokens

    merged = _merge_around(tokens, separators=("/",))
    if len(merged) == n_cols:
        return merged
    merged = _merge_around(merged, separators=("-",))
    if len(merged) == n_cols:
        return merged
    return None


def _merge_around(tokens: list[str], *, separators: tuple[str, ...]) -> list[str]:
    """Mescla padrões ``[A, SEP, B]`` em ``"A SEP B"`` (1 token só).

    Usado pra reconstruir cabeçalhos compostos depois do split por
    whitespace ter quebrado "Pré-Req / Co-req" em 3 pedaços.
    """
    out: list[str] = []
    i = 0
    while i < len(tokens):
        if i + 2 < len(tokens) and tokens[i + 1] in separators:
            out.append(f"{tokens[i]} {tokens[i + 1]} {tokens[i + 2]}")
            i += 3
        else:
            out.append(tokens[i])
            i += 1
    return out


def _clean_cell(cell) -> str:
    """Normaliza uma célula: None → '', múltiplos espaços/quebras → 1 espaço."""
    if cell is None:
        return ""
    return " ".join(str(cell).split())


def _clean_cell_keep_newlines(cell) -> str:
    """Igual a ``_clean_cell``, mas preserva ``\\n``. Usado pra detectar
    e parsear banner rows onde título e cabeçalho real estão na mesma
    célula separados por quebra de linha."""
    if cell is None:
        return ""
    s = re.sub(r"[ \t]+", " ", str(cell))
    # Tira espaços nas extremidades de cada linha sem perder o \n.
    return "\n".join(ln.strip() for ln in s.split("\n")).strip()


def _flat(s: str) -> str:
    """Achata espaços/quebras em string única separada por espaço simples."""
    return " ".join(s.split())


def _extract_docx(file_bytes: bytes) -> str:
    """
    Extrai texto de DOCX usando python-docx.
    Mantém a estrutura de parágrafos.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extrai também texto de tabelas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return _clean_text("\n\n".join(paragraphs))


def _clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Normaliza para que acentos no texto e na query se equivalham
    text = unicodedata.normalize("NFC", text)
    return text.strip()
