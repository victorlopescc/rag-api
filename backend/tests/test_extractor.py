"""Testa extração de texto para TXT/PDF/DOCX."""
import io

import docx
import fitz
import pytest

from pipeline.extractor import (
    _clean_cell,
    _extract_headers,
    _merge_around,
    _parse_header_line,
    _rect_overlaps,
    _render_table_as_sentences,
    extract_text,
)


def test_extract_txt_decodes_utf8():
    text = extract_text("Olá mundo — acentos".encode("utf-8"), "a.txt")
    assert "Olá" in text
    assert "mundo" in text


def test_extract_txt_with_bad_bytes_does_not_raise():
    # Byte inválido de UTF-8 — errors="replace" mantém o resto legível.
    text = extract_text(b"Hello \xff world", "a.txt")
    assert "Hello" in text
    assert "world" in text


def test_extract_unknown_extension_raises():
    with pytest.raises(ValueError):
        extract_text(b"x", "file.xyz")


def test_extract_pdf_returns_page_markers():
    pdf_bytes = _build_pdf(["Página 1 — regulamento", "Página 2 — estágio"])
    text = extract_text(pdf_bytes, "doc.pdf")
    assert "Página 1" in text
    assert "regulamento" in text
    assert "estágio" in text


def test_extract_pdf_skips_empty_pages():
    pdf_bytes = _build_pdf(["conteúdo", ""])
    text = extract_text(pdf_bytes, "doc.pdf")
    assert "conteúdo" in text


def test_extract_docx_reads_paragraphs_and_tables():
    docx_bytes = _build_docx(
        paragraphs=["Parágrafo A", "Parágrafo B"],
        table_rows=[["c1", "c2"], ["c3", "c4"]],
    )
    text = extract_text(docx_bytes, "doc.docx")
    assert "Parágrafo A" in text
    assert "c1 | c2" in text
    assert "c3 | c4" in text


def test_clean_text_normalizes_whitespace():
    # Muitos espaços e quebras — extrator reduz.
    raw = "a   b\n\n\n\n\nc"
    pdf = _build_pdf([raw])
    text = extract_text(pdf, "x.pdf")
    assert "   " not in text
    assert "\n\n\n" not in text


# --- Tabelas em PDF --------------------------------------------------------

def test_render_table_repeats_headers_in_each_row():
    """Cabeçalho deve aparecer em CADA linha de dados — esse é o ponto
    central do fix: chunker pode partir a tabela, mas todo chunk carrega
    contexto das colunas."""
    table = _FakeTable([
        ["Disciplina", "CH", "Período"],
        ["AEDs1", "120", "1"],
        ["TGC", "120", "5"],
    ])
    result = _render_table_as_sentences(table)
    assert "Disciplina: AEDs1" in result
    assert "CH: 120" in result
    assert "Período: 1" in result
    # Cabeçalho repetido na segunda linha
    assert "Disciplina: TGC" in result
    assert "Período: 5" in result


def test_render_table_skips_empty_rows():
    table = _FakeTable([
        ["A", "B"],
        ["x", "y"],
        ["", ""],  # totalmente vazia
        [None, None],  # idem
        ["w", "z"],
    ])
    result = _render_table_as_sentences(table)
    lines = [ln for ln in result.split("\n") if ln.strip()]
    assert len(lines) == 2
    assert "A: x" in result
    assert "A: w" in result


def test_render_table_handles_missing_header_cell():
    """Header com célula vazia recebe placeholder 'col_N' pra não gerar
    sentenças tipo ': 120' (sem chave)."""
    table = _FakeTable([
        ["Disciplina", "", "Período"],
        ["AEDs1", "120", "1"],
    ])
    result = _render_table_as_sentences(table)
    assert "Disciplina: AEDs1" in result
    assert "col_2: 120" in result
    assert "Período: 1" in result


def test_render_table_returns_empty_for_no_data_rows():
    # Só cabeçalho → não há nada de útil pra renderizar
    table = _FakeTable([["A", "B"]])
    assert _render_table_as_sentences(table) == ""


def test_render_table_returns_empty_when_extract_fails():
    table = _FakeTable(_raise=True)
    assert _render_table_as_sentences(table) == ""


# --- Banner row detection -------------------------------------------------

def test_extract_headers_normal_case_uses_row0():
    """Cabeçalho preenchido em ≥ metade das colunas → usado direto."""
    rows = [
        ["Disciplina", "CH", "Período"],
        ["AEDs1", "120", "1"],
    ]
    headers, data_start = _extract_headers(rows)
    assert headers == ["Disciplina", "CH", "Período"]
    assert data_start == 1


def test_extract_headers_partial_header_still_uses_row0():
    """Header com 1 célula vazia mas ≥ metade preenchida → ainda usa row 0."""
    rows = [
        ["Disciplina", "", "Período"],
        ["AEDs1", "120", "1"],
    ]
    headers, data_start = _extract_headers(rows)
    # Célula vazia vira col_N pra evitar header ambíguo.
    assert headers == ["Disciplina", "col_2", "Período"]
    assert data_start == 1


def test_extract_headers_banner_row_parses_real_headers_after_newline():
    """Banner row real do caso Grade Computação: título + headers
    na mesma célula separados por \\n. Deve parsear a última linha como
    cabeçalho."""
    banner = (
        "Curso de Ciência da Computação – Currículo 75902(PL)\n"
        "Per Disciplina CH Pré-Requisito / Co-requisito ID"
    )
    rows = [
        [banner, None, None, None, None],
        ["1", "Algoritmos e Estruturas de Dados I", "120", "---", "AEDs1"],
    ]
    headers, data_start = _extract_headers(rows)
    assert headers == [
        "Per",
        "Disciplina",
        "CH",
        "Pré-Requisito / Co-requisito",
        "ID",
    ]
    assert data_start == 1


def test_extract_headers_banner_without_embedded_headers_falls_back():
    """Banner com só o título (1 linha) → não conseguimos parsear, cai
    no fallback ``col_N`` em vez de gerar headers tóxicos."""
    rows = [
        ["Só um título sem cabeçalho embutido", None, None, None, None],
        ["1", "Algoritmos", "120", "---", "AEDs1"],
    ]
    headers, data_start = _extract_headers(rows)
    assert headers == ["col_1", "col_2", "col_3", "col_4", "col_5"]
    assert data_start == 1


def test_extract_headers_unparseable_count_falls_back_to_generic():
    """Banner com 2 linhas mas a 2ª não bate o número de colunas → fallback."""
    rows = [
        ["título\nA B", None, None, None, None],  # só 2 tokens mas 5 cols
        ["1", "x", "y", "z", "w"],
    ]
    headers, _ = _extract_headers(rows)
    assert headers == ["col_1", "col_2", "col_3", "col_4", "col_5"]


def test_render_table_with_banner_row_produces_clean_chunks():
    """Integração: tabela com banner row deve produzir sentenças limpas
    SEM o título do curso poluindo cada linha."""
    banner = (
        "Curso de Ciência da Computação – Currículo 75902(PL)\n"
        "Per Disciplina CH Pré-Requisito / Co-requisito ID"
    )
    table = _FakeTable([
        [banner, None, None, None, None],
        ["1", "Algoritmos e Estruturas de Dados I", "120", "---", "AEDs1"],
        ["2", "Arquitetura de Computadores I", "80", "AEDs1", "AC1"],
    ])
    result = _render_table_as_sentences(table)
    # Headers parseados, sem o título do curso vazando.
    assert "Per: 1" in result
    assert "Disciplina: Algoritmos e Estruturas de Dados I" in result
    assert "CH: 120" in result
    assert "Pré-Requisito / Co-requisito: ---" in result
    assert "ID: AEDs1" in result
    assert "Per: 2" in result
    assert "Disciplina: Arquitetura de Computadores I" in result
    # Crítico: o título do CURSO NÃO pode aparecer dentro das linhas.
    assert "Curso de Ciência" not in result


# --- parse_header_line / merge_around -------------------------------------

def test_parse_header_line_exact_match():
    assert _parse_header_line("A B C", 3) == ["A", "B", "C"]


def test_parse_header_line_merges_around_slash():
    out = _parse_header_line("Per Disciplina CH Pré-Req / Co-req ID", 5)
    assert out == ["Per", "Disciplina", "CH", "Pré-Req / Co-req", "ID"]


def test_parse_header_line_returns_none_when_no_match():
    # 5 tokens, queremos 3 — não dá pra mesclar com / nem -.
    assert _parse_header_line("A B C D E", 3) is None


def test_parse_header_line_empty_returns_none():
    assert _parse_header_line("", 3) is None
    assert _parse_header_line("   ", 3) is None


def test_merge_around_slash_collapses_triplet():
    out = _merge_around(["A", "/", "B", "C"], separators=("/",))
    assert out == ["A / B", "C"]


def test_merge_around_no_separator_returns_input():
    out = _merge_around(["A", "B", "C"], separators=("/",))
    assert out == ["A", "B", "C"]


def test_clean_cell_handles_none_and_whitespace():
    assert _clean_cell(None) == ""
    assert _clean_cell("  a   b\nc  ") == "a b c"
    assert _clean_cell(120) == "120"


def test_rect_overlaps_detects_majority_overlap():
    a = fitz.Rect(0, 0, 10, 10)        # área 100
    b = fitz.Rect(0, 0, 10, 10)        # idem
    assert _rect_overlaps(a, b)

    a = fitz.Rect(0, 0, 10, 10)
    b = fitz.Rect(20, 20, 30, 30)      # disjunto
    assert not _rect_overlaps(a, b)

    a = fitz.Rect(0, 0, 10, 10)        # área 100
    b = fitz.Rect(0, 0, 2, 2)           # interseção 4 → 4% < 50%
    assert not _rect_overlaps(a, b)


def test_extract_pdf_with_drawn_table_uses_structured_format():
    """Integração: PDF com grade desenhada deve produzir sentenças
    estruturadas em vez de células fragmentadas."""
    pdf_bytes = _build_pdf_with_table(
        headers=["Disciplina", "CH", "Periodo"],
        rows=[
            ["AEDs1", "120", "1"],
            ["TGC", "120", "5"],
        ],
    )
    text = extract_text(pdf_bytes, "grade.pdf")
    # A versão estruturada deve aparecer com cabeçalho colado ao valor
    assert "Disciplina: AEDs1" in text
    assert "CH: 120" in text
    assert "Periodo: 1" in text
    assert "Disciplina: TGC" in text


# --- helpers ---------------------------------------------------------------

def _build_pdf(pages: list[str]) -> bytes:
    doc = fitz.open()
    for content in pages:
        page = doc.new_page()
        if content:
            page.insert_text((72, 72), content)
    buf = doc.tobytes()
    doc.close()
    return buf


def _build_pdf_with_table(headers: list[str], rows: list[list[str]]) -> bytes:
    """Desenha uma grade real (linhas + texto) que o find_tables consegue
    detectar. PyMuPDF detecta tabelas a partir de vetores de linha + alinhamento
    de texto."""
    doc = fitz.open()
    page = doc.new_page()

    n_cols = len(headers)
    all_rows = [headers] + rows
    n_rows = len(all_rows)

    col_w = 120
    row_h = 25
    x0 = 50
    y0 = 100

    # Grade — linhas horizontais e verticais
    for i in range(n_rows + 1):
        y = y0 + i * row_h
        page.draw_line(
            fitz.Point(x0, y),
            fitz.Point(x0 + n_cols * col_w, y),
        )
    for j in range(n_cols + 1):
        x = x0 + j * col_w
        page.draw_line(
            fitz.Point(x, y0),
            fitz.Point(x, y0 + n_rows * row_h),
        )

    # Texto em cada célula
    for i, row in enumerate(all_rows):
        for j, cell in enumerate(row):
            x = x0 + j * col_w + 5
            y = y0 + i * row_h + 16
            page.insert_text(fitz.Point(x, y), str(cell))

    buf = doc.tobytes()
    doc.close()
    return buf


def _build_docx(paragraphs: list[str], table_rows: list[list[str]]) -> bytes:
    d = docx.Document()
    for p in paragraphs:
        d.add_paragraph(p)
    if table_rows:
        table = d.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row in enumerate(table_rows):
            for j, val in enumerate(row):
                table.cell(i, j).text = val
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class _FakeTable:
    """Imita ``fitz.Table``: tem só o método ``.extract()`` que o
    renderer chama. Permite testar a lógica de renderização sem depender
    de detecção via PyMuPDF (que requer um PDF real com grade)."""

    def __init__(self, rows: list[list] | None = None, _raise: bool = False):
        self._rows = rows or []
        self._raise = _raise

    def extract(self):
        if self._raise:
            raise RuntimeError("simulated extract failure")
        return self._rows
